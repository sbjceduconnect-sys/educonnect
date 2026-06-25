"""
Convert Markdown user guides to professionally formatted DOCX files.
"""
import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Theme colors ──────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x3F, 0x6B)
ORANGE = RGBColor(0xF0, 0x78, 0x30)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x33, 0x33, 0x33)
GRAY   = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
TIP_GREEN  = RGBColor(0x1A, 0x7F, 0x37)
WARN_AMBER = RGBColor(0xBF, 0x80, 0x00)
CAUT_RED   = RGBColor(0xCF, 0x22, 0x2E)
NOTE_BLUE  = RGBColor(0x0D, 0x6E, 0xFD)
IMP_PURPLE = RGBColor(0x88, 0x57, 0xD0)

BG_TIP   = "D1FAE5"
BG_WARN  = "FEF3C7"
BG_CAUT  = "FEE2E2"
BG_NOTE  = "DBEAFE"
BG_IMP   = "EDE9FE"


def setup_styles(doc):
    """Create custom styles for the document."""
    style = doc.styles

    # ── Normal ────────────────────────────────────────────────────
    normal = style['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # ── Heading 1 ─────────────────────────────────────────────────
    h1 = style['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)

    # ── Heading 2 ─────────────────────────────────────────────────
    h2 = style['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = NAVY
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    # ── Heading 3 ─────────────────────────────────────────────────
    h3 = style['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = ORANGE
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    # ── Code style ────────────────────────────────────────────────
    if 'Code Block' not in [s.name for s in style]:
        code_style = style.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(9.5)
        code_style.font.color.rgb = DARK
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(4)
        code_style.paragraph_format.left_indent = Cm(0.5)

    # ── TOC Entry style ───────────────────────────────────────────
    if 'TOC Entry' not in [s.name for s in style]:
        toc_style = style.add_style('TOC Entry', WD_STYLE_TYPE.PARAGRAPH)
        toc_style.font.name = 'Calibri'
        toc_style.font.size = Pt(11)
        toc_style.font.color.rgb = NAVY
        toc_style.paragraph_format.space_after = Pt(2)
        toc_style.paragraph_format.left_indent = Cm(0.5)


def add_title_page(doc, title, subtitle, emoji):
    """Create a professional title/cover page."""
    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Emoji icon
    p_emoji = doc.add_paragraph()
    p_emoji.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_emoji.add_run(emoji)
    run.font.size = Pt(48)

    doc.add_paragraph()

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = ORANGE
    run.font.name = 'Calibri'
    run.font.bold = True

    doc.add_paragraph()

    # Divider line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line.add_run("━" * 40)
    run.font.color.rgb = ORANGE
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Institution
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_inst.add_run("SaiBalaji Junior College")
    run.font.size = Pt(16)
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'

    # College Management System
    p_cms = doc.add_paragraph()
    p_cms.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cms.add_run("EduConnect — College Management System")
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'

    doc.add_paragraph()

    # Version
    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_ver.add_run("Version 1.0  •  June 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = LIGHT_GRAY
    run.font.name = 'Calibri'

    # Page break
    doc.add_page_break()


def set_cell_shading(cell, color_hex):
    """Apply background color to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        set_cell_shading(cell, "1B3F6B")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            # Parse bold markers
            parts = re.split(r'\*\*(.+?)\*\*', cell_text)
            for j, part in enumerate(parts):
                run = p.add_run(part)
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                if j % 2 == 1:
                    run.bold = True
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0F4F8")

    doc.add_paragraph()  # spacing after table


def add_alert_box(doc, alert_type, text):
    """Add a colored callout/alert box."""
    labels = {
        "NOTE": ("📝 NOTE", BG_NOTE, NOTE_BLUE),
        "TIP": ("💡 TIP", BG_TIP, TIP_GREEN),
        "IMPORTANT": ("⚠️ IMPORTANT", BG_IMP, IMP_PURPLE),
        "WARNING": ("⚠️ WARNING", BG_WARN, WARN_AMBER),
        "CAUTION": ("🛑 CAUTION", BG_CAUT, CAUT_RED),
    }
    label, bg_color, text_color = labels.get(alert_type, ("NOTE", BG_NOTE, NOTE_BLUE))

    # Create a single-cell table for the alert box
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg_color)

    # Label
    p_label = cell.paragraphs[0]
    run = p_label.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = text_color
    run.font.name = 'Calibri'

    # Content
    p_content = cell.add_paragraph()
    # Parse bold text
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for j, part in enumerate(parts):
        run = p_content.add_run(part)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = text_color
        if j % 2 == 1:
            run.bold = True

    # Set border
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="0" w:color="{text_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    doc.add_paragraph()


def add_formatted_run(paragraph, text):
    """Add a run with inline bold/code/emoji formatting."""
    # Split by bold markers and backtick code
    # Process bold first
    bold_parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(bold_parts):
        if i % 2 == 1:
            # Bold text
            code_parts = re.split(r'`(.+?)`', part)
            for j, cp in enumerate(code_parts):
                run = paragraph.add_run(cp)
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                if j % 2 == 1:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
        else:
            # Normal text with possible inline code
            code_parts = re.split(r'`(.+?)`', part)
            for j, cp in enumerate(code_parts):
                run = paragraph.add_run(cp)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                if j % 2 == 1:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)


def parse_md_to_docx(md_path, docx_path, title, subtitle, emoji):
    """Parse markdown and create a formatted DOCX file."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    setup_styles(doc)
    add_title_page(doc, title, subtitle, emoji)

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_headers = []
    table_rows = []
    in_alert = False
    alert_type = ""
    alert_lines = []

    while i < len(lines):
        line = lines[i].rstrip('\n').rstrip('\r')

        # ── Skip the first title line & metadata lines ────────────
        if i < 5 and (line.startswith('# ') or line.startswith('**SaiBalaji') or line.startswith('*Version')):
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────
        if line.strip() == '---':
            p = doc.add_paragraph()
            run = p.add_run("━" * 60)
            run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
            run.font.size = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # ── Code blocks ───────────────────────────────────────────
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                for cl in code_lines:
                    p = doc.add_paragraph(cl, style='Code Block')
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Alert boxes (GitHub style) ────────────────────────────
        alert_match = re.match(r'^>\s*\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]', line)
        if alert_match:
            alert_type = alert_match.group(1)
            in_alert = True
            alert_lines = []
            i += 1
            continue

        if in_alert:
            if line.startswith('> '):
                alert_lines.append(line[2:].strip())
                i += 1
                continue
            elif line.strip() == '':
                # End of alert
                add_alert_box(doc, alert_type, ' '.join(alert_lines))
                in_alert = False
                alert_lines = []
                i += 1
                continue
            else:
                # End of alert (non-empty non-quote line)
                add_alert_box(doc, alert_type, ' '.join(alert_lines))
                in_alert = False
                alert_lines = []
                # Don't increment i, process this line normally

        # ── Tables ────────────────────────────────────────────────
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            if not in_table:
                # Check if next line is separator
                if i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i + 1].strip()):
                    table_headers = cells
                    in_table = True
                    i += 2  # skip header and separator
                    continue
                else:
                    # Could be a data row continuation
                    table_rows.append(cells)
                    i += 1
                    continue
            else:
                table_rows.append(cells)
                i += 1
                continue

        if in_table and (not line.strip().startswith('|') or line.strip() == ''):
            # End of table
            add_table(doc, table_headers, table_rows)
            table_headers = []
            table_rows = []
            in_table = False
            if line.strip() == '':
                i += 1
                continue
            # Process current line normally

        # ── Headings ──────────────────────────────────────────────
        h1_match = re.match(r'^## (.+)', line)
        h2_match = re.match(r'^### (.+)', line)

        if h1_match:
            text = h1_match.group(1).strip()
            doc.add_heading(text, level=1)
            i += 1
            continue

        if h2_match:
            text = h2_match.group(1).strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        # ── Table of Contents entries ─────────────────────────────
        toc_match = re.match(r'^(\d+)\.\s+\[(.+?)\]', line)
        if toc_match:
            num = toc_match.group(1)
            text = toc_match.group(2)
            p = doc.add_paragraph(style='TOC Entry')
            run = p.add_run(f"{num}. {text}")
            run.font.size = Pt(11)
            run.font.color.rgb = NAVY
            i += 1
            continue

        # ── Numbered lists ────────────────────────────────────────
        num_match = re.match(r'^(\d+)\.\s+(.+)', line)
        if num_match:
            num = num_match.group(1)
            text = num_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(f"{num}. ")
            run.bold = True
            run.font.color.rgb = ORANGE
            run.font.name = 'Calibri'
            add_formatted_run(p, text)
            i += 1
            continue

        # ── Bullet lists ──────────────────────────────────────────
        bullet_match = re.match(r'^(\s*)-\s+(.+)', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(2)
            p = doc.add_paragraph()
            indent_cm = 1.0 + (indent / 2) * 0.5
            p.paragraph_format.left_indent = Cm(indent_cm)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run("● ")
            run.font.color.rgb = ORANGE
            run.font.size = Pt(9)
            add_formatted_run(p, text)
            i += 1
            continue

        # ── Empty lines ───────────────────────────────────────────
        if line.strip() == '':
            i += 1
            continue

        # ── Regular paragraphs ────────────────────────────────────
        p = doc.add_paragraph()
        add_formatted_run(p, line)
        i += 1

    # Handle any remaining table/alert
    if in_table and table_headers:
        add_table(doc, table_headers, table_rows)
    if in_alert and alert_lines:
        add_alert_box(doc, alert_type, ' '.join(alert_lines))

    # ── Footer ────────────────────────────────────────────────────
    doc.add_page_break()
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(6):
        doc.add_paragraph()
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_end.add_run("━" * 30)
    run.font.color.rgb = ORANGE
    run.font.size = Pt(10)

    p_end2 = doc.add_paragraph()
    p_end2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_end2.add_run("EduConnect © 2026 — SaiBalaji Junior College")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'

    p_end3 = doc.add_paragraph()
    p_end3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_end3.add_run("All rights reserved.")
    run.font.size = Pt(9)
    run.font.color.rgb = LIGHT_GRAY
    run.font.name = 'Calibri'

    doc.save(docx_path)
    print(f"  ✅ Created: {docx_path}")


# ── Main ──────────────────────────────────────────────────────────────
if __name__ == '__main__':
    docs_dir = os.path.dirname(os.path.abspath(__file__))

    guides = [
        {
            "md": os.path.join(docs_dir, "Admin_User_Guide.md"),
            "docx": os.path.join(docs_dir, "Admin_User_Guide.docx"),
            "title": "Administrator\nUser Guide",
            "subtitle": "Admin Panel",
            "emoji": "📘",
        },
        {
            "md": os.path.join(docs_dir, "Teacher_User_Guide.md"),
            "docx": os.path.join(docs_dir, "Teacher_User_Guide.docx"),
            "title": "Teacher\nUser Guide",
            "subtitle": "Teaching Portal",
            "emoji": "📗",
        },
        {
            "md": os.path.join(docs_dir, "Student_User_Guide.md"),
            "docx": os.path.join(docs_dir, "Student_User_Guide.docx"),
            "title": "Student\nUser Guide",
            "subtitle": "Student Portal",
            "emoji": "📙",
        },
    ]

    print("\n🔄 Converting Markdown to DOCX...\n")
    for g in guides:
        if os.path.exists(g["md"]):
            parse_md_to_docx(g["md"], g["docx"], g["title"], g["subtitle"], g["emoji"])
        else:
            print(f"  ❌ Not found: {g['md']}")

    print("\n✅ All done!\n")
